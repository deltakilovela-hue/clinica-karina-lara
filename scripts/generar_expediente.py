#!/usr/bin/env python3
"""
Genera Expediente_Clinico_Nutricional.docx con datos reales del paciente.
Uso: python3 generar_expediente.py <input.json> <output.docx>
"""

import sys, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# ─── Argumentos ─────────────────────────────────────────────────────────────
if len(sys.argv) < 3:
    print("Uso: generar_expediente.py <input.json> <output.docx>", file=sys.stderr)
    sys.exit(1)

with open(sys.argv[1], encoding='utf-8') as f:
    datos = json.load(f)

output_path = sys.argv[2]

# ─── Datos del paciente ──────────────────────────────────────────────────────
pac     = datos.get('paciente', {}) or {}
historia = datos.get('historia', {}) or {}
mediciones = datos.get('mediciones', []) or []
diagnostico = datos.get('diagnostico', {}) or {}
planes  = datos.get('planes', []) or []
seguimientos = datos.get('seguimientos', []) or []

def v(d, *keys, default='—'):
    """Obtiene un valor de un dict de manera segura."""
    for k in keys:
        if isinstance(d, dict):
            d = d.get(k)
        else:
            return default
    if d is None or d == '':
        return default
    if isinstance(d, list):
        return ', '.join(str(x) for x in d) if d else default
    return str(d)

def fmt_fecha(ts_str):
    """Convierte timestamp de Firebase (dict con _seconds) o string a fecha legible."""
    if not ts_str:
        return '—'
    if isinstance(ts_str, dict):
        try:
            from datetime import timezone
            dt = datetime.fromtimestamp(ts_str.get('_seconds', 0), tz=timezone.utc)
            return dt.strftime('%d/%m/%Y')
        except:
            return '—'
    if isinstance(ts_str, str):
        return ts_str[:10]
    return str(ts_str)

# ─── Documento ───────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ─── Colores ─────────────────────────────────────────────────────────────────
ROJO    = RGBColor(0x7B, 0x1B, 0x2A)
VERDE   = RGBColor(0x2D, 0x6A, 0x4F)
GRIS    = RGBColor(0x55, 0x55, 0x55)
NEGRO   = RGBColor(0x1A, 0x1A, 0x1A)
DORADO  = RGBColor(0x8B, 0x69, 0x14)
BLANCO_H= "7B1B2A"

# ─── Helpers ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color='CCCCCC'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading_para(text, color=ROJO, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.bold  = True
    run.font.size  = Pt(size)
    run.font.name  = 'Arial'
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), str(color))
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def cell_label(cell, text, size=9, bold=True, color=DORADO):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Arial'
    cell.paragraphs[0].paragraph_format.space_after = Pt(1)

def cell_value(cell, text, size=10, bold=False, wrap=False):
    if len(cell.paragraphs) > 0:
        p = cell.paragraphs[-1]
    else:
        p = cell.add_paragraph()
    run = p.add_run(str(text) if text else '—')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Arial'
    run.font.color.rgb = NEGRO

def label_value_cell(cell, label, value, label_size=9, value_size=10):
    cell.paragraphs[0].clear()
    r1 = cell.paragraphs[0].add_run(label + '\n')
    r1.font.bold = True
    r1.font.size = Pt(label_size)
    r1.font.color.rgb = DORADO
    r1.font.name = 'Arial'
    r2 = cell.paragraphs[0].add_run(str(value) if value else '—')
    r2.font.size = Pt(value_size)
    r2.font.name = 'Arial'
    r2.font.color.rgb = NEGRO

def add_table(rows, cols, col_widths=None):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    if col_widths:
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = Cm(col_widths[i])
    return tbl

# ════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════════════════════════════════════════
# Header rojo
p_header = doc.add_paragraph()
p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_header.paragraph_format.space_before = Pt(0)
p_header.paragraph_format.space_after  = Pt(4)
r = p_header.add_run('CLÍNICA KARINA LARA · NUTRICIÓN PEDIÁTRICA')
r.font.bold = True; r.font.size = Pt(14); r.font.name = 'Arial'
r.font.color.rgb = ROJO

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(2)
r2 = p_sub.add_run('EXPEDIENTE CLÍNICO NUTRICIONAL')
r2.font.bold = True; r2.font.size = Pt(13); r2.font.name = 'Arial'
r2.font.color.rgb = NEGRO

p_fecha = doc.add_paragraph()
p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fecha.paragraph_format.space_after = Pt(10)
r3 = p_fecha.add_run(f'Generado: {datetime.now().strftime("%d/%m/%Y %H:%M")}  ·  Lic. Karina Lara')
r3.font.size = Pt(10); r3.font.name = 'Arial'; r3.font.color.rgb = GRIS

# Línea separadora
p_line = doc.add_paragraph()
p_line.paragraph_format.space_after = Pt(6)
pPr = p_line._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom')
bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '7B1B2A')
pBdr.append(bot); pPr.append(pBdr)

# ════════════════════════════════════════════════════════════════════════════
#  1. DATOS GENERALES
# ════════════════════════════════════════════════════════════════════════════
heading_para('1. DATOS GENERALES DEL PACIENTE')

nombre = v(pac, 'nombre')
edad = v(pac, 'edad')
sexo = v(pac, 'sexo')
tutor = v(pac, 'tutor')
telefono = v(pac, 'telefono')
correo = v(pac, 'correo')
fecha_nac = fmt_fecha(pac.get('fechaNacimiento'))
motivo = v(pac, 'motivoConsulta')

tbl1 = add_table(4, 4, [4, 4, 4, 4])
datos_gen = [
    [('Nombre del Paciente', nombre), ('Edad', edad + ' años' if edad != '—' else '—'),
     ('Sexo', sexo), ('Fecha de Nacimiento', fecha_nac)],
    [('Tutor / Padre / Madre', tutor), ('Teléfono de Contacto', telefono),
     ('Correo Electrónico', correo), ('No. Expediente', v(pac, 'id'))],
    [('Motivo de Consulta', motivo), None, None, None],
    [('Alergias Conocidas', v(historia, 'alergias')), None,
     ('Intolerancias Conocidas', v(historia, 'intolerancias')), None],
]

for ri, row_data in enumerate(datos_gen):
    row = tbl1.rows[ri]
    ci = 0
    for item in row_data:
        if item is None:
            ci += 1
            continue
        label, val = item
        label_value_cell(row.cells[ci], label, val)
        set_cell_bg(row.cells[ci], 'FAF7F2' if ri % 2 == 0 else 'FFFFFF')
        set_cell_border(row.cells[ci])
        ci += 1

# Merge celda motivo (ocupa toda la fila)
tbl1.rows[2].cells[0].merge(tbl1.rows[2].cells[3])
tbl1.rows[3].cells[0].merge(tbl1.rows[3].cells[1])
tbl1.rows[3].cells[2].merge(tbl1.rows[3].cells[3])

# ════════════════════════════════════════════════════════════════════════════
#  2. ANTECEDENTES PERINATALES Y MÉDICOS
# ════════════════════════════════════════════════════════════════════════════
heading_para('2. ANTECEDENTES PERINATALES Y MÉDICOS')

tbl2 = add_table(2, 4, [4, 4, 4, 4])
ant_data = [
    [('Tipo de Parto', v(historia, 'tipoParto')),
     ('Edad Gestacional', v(historia, 'edadGestacional')),
     ('Peso al Nacer', v(historia, 'pesoNacer')),
     ('Talla al Nacer', v(historia, 'tallaNacer'))],
    [('Lactancia Materna', v(historia, 'lactanciaMaterna')),
     ('Duración LM', v(historia, 'duracionLactancia')),
     ('Alimentación Complementaria', v(historia, 'alimentacionComplementaria')),
     ('Antecedentes Perinatales', v(historia, 'antecedentesPrenatales'))],
]
for ri, row_data in enumerate(ant_data):
    row = tbl2.rows[ri]
    for ci, (label, val) in enumerate(row_data):
        label_value_cell(row.cells[ci], label, val)
        set_cell_bg(row.cells[ci], 'FAF7F2' if ri % 2 == 0 else 'FFFFFF')
        set_cell_border(row.cells[ci])

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Diagnósticos Médicos Previos')
r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = DORADO; r.font.name = 'Arial'

tbl2b = add_table(1, 1)
c = tbl2b.rows[0].cells[0]
label_value_cell(c, 'Descripción', v(historia, 'diagnosticosPrevios'), value_size=10)
set_cell_border(c)

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(8)
p2.paragraph_format.space_after = Pt(4)
r2 = p2.add_run('Medicamentos Actuales')
r2.font.bold = True; r2.font.size = Pt(10); r2.font.color.rgb = DORADO; r2.font.name = 'Arial'

tbl2c = add_table(1, 4, [4.25, 4.25, 4.25, 3.25])
headers_med = ['Medicamento', 'Dosis', 'Frecuencia', 'Motivo']
for ci, h in enumerate(headers_med):
    cell = tbl2c.rows[0].cells[ci]
    label_value_cell(cell, h, v(historia, 'medicamentos') if ci == 0 else '—')
    set_cell_bg(cell, 'F5E8EB')
    set_cell_border(cell)

# ════════════════════════════════════════════════════════════════════════════
#  3. SÍNTOMAS GASTROINTESTINALES
# ════════════════════════════════════════════════════════════════════════════
heading_para('3. SÍNTOMAS GASTROINTESTINALES')

TODOS_SINTOMAS = [
    'Estreñimiento', 'Diarrea', 'Reflujo / Regurgitación', 'Distensión abdominal',
    'Flatulencias / Gases', 'Dolor abdominal', 'Náuseas', 'Vómito', 'Heces con moco',
    'Heces con sangre', 'Urgencia fecal', 'Incontinencia fecal', 'Sensación de vaciado incompleto',
    'Deposiciones muy frecuentes (+3/día)', 'Inapetencia / Falta de apetito',
    'Hiperfagia / Comer en exceso', 'Disfagia / Dificultad para tragar',
    'Arcadas al comer', 'Selectividad severa por texturas', 'Vómito emocional / conductual',
    'Cólico del lactante'
]

sintomas_paciente = historia.get('sintomasGI', []) or []
if isinstance(sintomas_paciente, str):
    sintomas_paciente = [sintomas_paciente] if sintomas_paciente else []
sintomas_otros = v(historia, 'sintomasGIOtros', default='')

tbl3 = add_table(1 + len(TODOS_SINTOMAS), 4, [6.5, 2.5, 2.5, 4.5])
headers3 = ['Síntoma', 'Presente', 'Intensidad (1–5)', 'Observaciones']
for ci, h in enumerate(headers3):
    cell = tbl3.rows[0].cells[ci]
    cell.paragraphs[0].clear()
    r = cell.paragraphs[0].add_run(h)
    r.font.bold = True; r.font.size = Pt(9); r.font.name = 'Arial'; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_bg(cell, '7B1B2A')

for ri, sintoma in enumerate(TODOS_SINTOMAS):
    row = tbl3.rows[ri + 1]
    presente = '✓' if sintoma in sintomas_paciente else '✗'
    color_bg = 'FFE8EC' if presente == '✓' else 'FFFFFF'

    cell_sn = row.cells[0]
    cell_sn.paragraphs[0].clear()
    r = cell_sn.paragraphs[0].add_run(sintoma)
    r.font.size = Pt(9); r.font.name = 'Arial'
    r.font.color.rgb = RGBColor(0x7B, 0x1B, 0x2A) if presente == '✓' else NEGRO
    r.font.bold = presente == '✓'
    set_cell_bg(cell_sn, color_bg)
    set_cell_border(cell_sn)

    cell_pr = row.cells[1]
    cell_pr.paragraphs[0].clear()
    r2 = cell_pr.paragraphs[0].add_run(presente)
    r2.font.size = Pt(10); r2.font.bold = True; r2.font.name = 'Arial'
    r2.font.color.rgb = RGBColor(0x7B, 0x1B, 0x2A) if presente == '✓' else GRIS
    cell_pr.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cell_pr, color_bg)
    set_cell_border(cell_pr)

    for ci in [2, 3]:
        set_cell_bg(row.cells[ci], color_bg)
        set_cell_border(row.cells[ci])

if sintomas_otros and sintomas_otros != '—':
    p_otros = doc.add_paragraph()
    p_otros.paragraph_format.space_before = Pt(4)
    r_o = p_otros.add_run(f'Otros síntomas: {sintomas_otros}')
    r_o.font.size = Pt(10); r_o.font.italic = True; r_o.font.name = 'Arial'; r_o.font.color.rgb = GRIS

# ════════════════════════════════════════════════════════════════════════════
#  4. ANTROPOMETRÍA — HISTORIAL
# ════════════════════════════════════════════════════════════════════════════
heading_para('4. ANTROPOMETRÍA — HISTORIAL DE MEDICIONES')

headers4 = ['Fecha', 'Peso (kg)', 'Talla (cm)', 'IMC', 'P. Peso/E', 'P. Talla/E',
            'Interp. Peso', 'Interp. Talla', 'Observaciones']
col_w4 = [2.5, 1.8, 1.8, 1.5, 1.7, 1.7, 2.5, 2.5, 2.0]
n_meds = max(len(mediciones), 1)
tbl4 = add_table(1 + n_meds, 9, col_w4)

for ci, h in enumerate(headers4):
    cell = tbl4.rows[0].cells[ci]
    cell.paragraphs[0].clear()
    r = cell.paragraphs[0].add_run(h)
    r.font.bold = True; r.font.size = Pt(8.5); r.font.name = 'Arial'; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    set_cell_bg(cell, '37474F')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for ri, med in enumerate(mediciones if mediciones else [{}]):
    row = tbl4.rows[ri + 1]
    bg = 'FAF7F2' if ri % 2 == 0 else 'FFFFFF'
    vals = [
        fmt_fecha(med.get('fechaCreacion') or med.get('fecha')),
        v(med, 'peso'), v(med, 'talla'), v(med, 'imc'),
        f"P{v(med, 'percentilPeso')}", f"P{v(med, 'percentilTalla')}",
        v(med, 'interpretacionPeso'), v(med, 'interpretacionTalla'), ''
    ]
    for ci, val in enumerate(vals):
        cell = row.cells[ci]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(val)
        r.font.size = Pt(9); r.font.name = 'Arial'; r.font.color.rgb = NEGRO
        set_cell_bg(cell, bg)
        set_cell_border(cell)

# ════════════════════════════════════════════════════════════════════════════
#  5. CONDUCTA Y SELECTIVIDAD ALIMENTARIA
# ════════════════════════════════════════════════════════════════════════════
heading_para('5. CONDUCTA Y SELECTIVIDAD ALIMENTARIA')

tbl5 = add_table(5, 2, [5.5, 10.5])
campos5 = [
    ('Conducta Alimentaria General', v(historia, 'conductaAlimentaria')),
    ('Alimentos Favoritos / Aceptados', v(historia, 'alimentosFavoritos')),
    ('Alimentos Rechazados', v(historia, 'alimentosRechazados')),
    ('Texturas Aceptadas', v(historia, 'texturasAceptadas')),
    ('Texturas Rechazadas', v(historia, 'texturasRechazadas')),
]
for ri, (label, val) in enumerate(campos5):
    row = tbl5.rows[ri]
    bg = 'FAF7F2' if ri % 2 == 0 else 'FFFFFF'
    cell_l = row.cells[0]; cell_v = row.cells[1]
    label_value_cell(cell_l, label, '')
    cell_l.paragraphs[0].runs[-1].text = ''  # clear default value
    r_v = cell_v.paragraphs[0].add_run(val)
    r_v.font.size = Pt(10); r_v.font.name = 'Arial'; r_v.font.color.rgb = NEGRO
    for c in [cell_l, cell_v]:
        set_cell_bg(c, bg); set_cell_border(c)

# ════════════════════════════════════════════════════════════════════════════
#  6. HÁBITOS DE VIDA
# ════════════════════════════════════════════════════════════════════════════
heading_para('6. HÁBITOS DE VIDA')

tbl6 = add_table(6, 3, [5, 7, 4])
headers6 = ['Aspecto', 'Descripción del Paciente', 'Observación Clínica']
for ci, h in enumerate(headers6):
    cell = tbl6.rows[0].cells[ci]
    cell.paragraphs[0].clear()
    r = cell.paragraphs[0].add_run(h)
    r.font.bold = True; r.font.size = Pt(9); r.font.name = 'Arial'; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    set_cell_bg(cell, BLANCO_H)

campos6 = [
    ('Horas de Sueño / Descanso', v(historia, 'horasSueno')),
    ('Actividad Física', v(historia, 'actividadFisica')),
    ('Horas de Pantalla / Tecnología', v(historia, 'horasPantalla')),
    ('Consumo de Agua (ml/día)', v(historia, 'consumoAgua')),
    ('Ambiente en las Comidas', v(historia, 'ambienteComidas')),
]
for ri, (aspecto, desc) in enumerate(campos6):
    row = tbl6.rows[ri + 1]
    bg = 'FAF7F2' if ri % 2 == 0 else 'FFFFFF'
    cell_l = row.cells[0]
    cell_l.paragraphs[0].clear()
    r = cell_l.paragraphs[0].add_run(aspecto)
    r.font.bold = True; r.font.size = Pt(9); r.font.name = 'Arial'; r.font.color.rgb = DORADO

    cell_d = row.cells[1]
    r_d = cell_d.paragraphs[0].add_run(desc)
    r_d.font.size = Pt(10); r_d.font.name = 'Arial'; r_d.font.color.rgb = NEGRO

    for c in [cell_l, cell_d, row.cells[2]]:
        set_cell_bg(c, bg); set_cell_border(c)

# ════════════════════════════════════════════════════════════════════════════
#  7. CONTEXTO FAMILIAR Y SOCIOECONÓMICO
# ════════════════════════════════════════════════════════════════════════════
heading_para('7. CONTEXTO FAMILIAR Y SOCIOECONÓMICO')

tbl7 = add_table(5, 2, [6, 10])
campos7 = [
    ('Tipo de Familia / Estructura', v(historia, 'tipoFamilia')),
    ('Nivel Socioeconómico', v(historia, 'nivelSocioeconomico')),
    ('Presupuesto Alimentario Semanal', v(historia, 'presupuestoAlimentario')),
    ('Persona que Prepara Alimentos', v(historia, 'personaPrepara')),
    ('Dinámica Familiar en la Alimentación', v(historia, 'dinamicaFamiliar')),
]
for ri, (label, val) in enumerate(campos7):
    row = tbl7.rows[ri]
    bg = 'FAF7F2' if ri % 2 == 0 else 'FFFFFF'
    cell_l = row.cells[0]
    cell_l.paragraphs[0].clear()
    r = cell_l.paragraphs[0].add_run(label)
    r.font.bold = True; r.font.size = Pt(9.5); r.font.name = 'Arial'; r.font.color.rgb = DORADO

    cell_v = row.cells[1]
    r_v = cell_v.paragraphs[0].add_run(val)
    r_v.font.size = Pt(10); r_v.font.name = 'Arial'; r_v.font.color.rgb = NEGRO

    for c in [cell_l, cell_v]:
        set_cell_bg(c, bg); set_cell_border(c)

# ════════════════════════════════════════════════════════════════════════════
#  8. DIAGNÓSTICO NUTRICIONAL
# ════════════════════════════════════════════════════════════════════════════
heading_para('8. DIAGNÓSTICO NUTRICIONAL')

diag_nutri = diagnostico.get('diagnosticoNutricional', []) or []
if isinstance(diag_nutri, str):
    diag_nutri = [diag_nutri] if diag_nutri else []
diag_texto = ', '.join(diag_nutri) if diag_nutri else '—'

diag_fecha = fmt_fecha(diagnostico.get('fechaCreacion'))

tbl8 = add_table(4, 2, [7, 9])
diag_campos = [
    ('Diagnóstico Nutricional', diag_texto),
    ('Fecha del Diagnóstico', diag_fecha),
    ('IMC al momento del Dx', v(mediciones[0] if mediciones else {}, 'imc')),
    ('Observaciones Clínicas', v(diagnostico, 'observaciones')),
]
for ri, (label, val) in enumerate(diag_campos):
    row = tbl8.rows[ri]
    bg = 'FAF7F2' if ri % 2 == 0 else 'FFFFFF'
    cell_l = row.cells[0]
    cell_l.paragraphs[0].clear()
    r = cell_l.paragraphs[0].add_run(label)
    r.font.bold = True; r.font.size = Pt(9.5); r.font.name = 'Arial'; r.font.color.rgb = DORADO

    cell_v = row.cells[1]
    r_v = cell_v.paragraphs[0].add_run(val)
    r_v.font.size = Pt(10); r_v.font.name = 'Arial'; r_v.font.color.rgb = NEGRO

    for c in [cell_l, cell_v]:
        set_cell_bg(c, bg); set_cell_border(c)

# ════════════════════════════════════════════════════════════════════════════
#  9. PLAN NUTRICIONAL VIGENTE
# ════════════════════════════════════════════════════════════════════════════
heading_para('9. PLAN NUTRICIONAL VIGENTE')

if planes:
    ultimo_plan = planes[0]
    plan_texto = str(ultimo_plan.get('texto', '—') or '—')
    plan_fecha = fmt_fecha(ultimo_plan.get('fechaCreacion'))

    p_pf = doc.add_paragraph()
    r_pf = p_pf.add_run(f'Generado el: {plan_fecha}')
    r_pf.font.size = Pt(10); r_pf.font.italic = True; r_pf.font.name = 'Arial'; r_pf.font.color.rgb = GRIS

    # Extraer sección del plan (hasta lista del súper)
    def extraer(texto, desde, hasta=None):
        i = texto.find(desde)
        if i == -1: return texto[:800] if len(texto) > 800 else texto
        start = i + len(desde)
        if not hasta: return texto[start:].strip()
        fin = texto.find(hasta, start)
        return texto[start:fin].strip() if fin != -1 else texto[start:].strip()

    plan_sec = extraer(plan_texto, '## PLAN NUTRICIONAL', '### LISTA DEL SÚPER')
    lista_sec = extraer(plan_texto, '### LISTA DEL SÚPER', '### PRESUPUESTO')

    # Limitar longitud para no crear doc enorme
    plan_muestra = plan_sec[:3000] if len(plan_sec) > 3000 else plan_sec

    tbl9 = add_table(1, 1)
    c9 = tbl9.rows[0].cells[0]
    c9.paragraphs[0].clear()
    r9 = c9.paragraphs[0].add_run(plan_muestra if plan_muestra else plan_texto[:2000])
    r9.font.size = Pt(9.5); r9.font.name = 'Arial'; r9.font.color.rgb = NEGRO
    set_cell_bg(c9, 'FAF7F2'); set_cell_border(c9)

    if lista_sec:
        p_ls = doc.add_paragraph()
        p_ls.paragraph_format.space_before = Pt(10)
        r_ls = p_ls.add_run('🛒 Lista del Súper')
        r_ls.font.bold = True; r_ls.font.size = Pt(11); r_ls.font.name = 'Arial'; r_ls.font.color.rgb = VERDE

        tbl9b = add_table(1, 1)
        c9b = tbl9b.rows[0].cells[0]
        c9b.paragraphs[0].clear()
        r9b = c9b.paragraphs[0].add_run(lista_sec[:2000])
        r9b.font.size = Pt(9.5); r9b.font.name = 'Arial'; r9b.font.color.rgb = NEGRO
        set_cell_bg(c9b, 'D8F3DC'); set_cell_border(c9b, 'B7E4C7')
else:
    p_no = doc.add_paragraph()
    r_no = p_no.add_run('Sin plan nutricional generado.')
    r_no.font.size = Pt(10); r_no.font.italic = True; r_no.font.name = 'Arial'; r_no.font.color.rgb = GRIS

# ════════════════════════════════════════════════════════════════════════════
#  10. SEGUIMIENTO EVOLUTIVO
# ════════════════════════════════════════════════════════════════════════════
heading_para('10. SEGUIMIENTO EVOLUTIVO')

BRISTOL_LABEL = {1:'Tipo 1 - Bolitas duras', 2:'Tipo 2 - Grumosa', 3:'Tipo 3 - Grietas',
                 4:'Tipo 4 - Lisa/suave ✓', 5:'Tipo 5 - Blanda', 6:'Tipo 6 - Semilíquida', 7:'Tipo 7 - Líquida'}

headers10 = ['Fecha', 'Peso (kg)', 'Bristol', 'Evacua./día', 'Tolerancia', 'Síntomas Activos', 'Notas de Evolución']
col_w10 = [2.2, 1.8, 3.0, 1.8, 2.5, 3.5, 4.2]
n_segs = max(len(seguimientos), 1)
tbl10 = add_table(1 + n_segs, 7, col_w10)

for ci, h in enumerate(headers10):
    cell = tbl10.rows[0].cells[ci]
    cell.paragraphs[0].clear()
    r = cell.paragraphs[0].add_run(h)
    r.font.bold = True; r.font.size = Pt(8.5); r.font.name = 'Arial'; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    set_cell_bg(cell, '37474F')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for ri, seg in enumerate(seguimientos if seguimientos else [{}]):
    row = tbl10.rows[ri + 1]
    bg = 'FAF7F2' if ri % 2 == 0 else 'FFFFFF'

    bristol_num = seg.get('bristolTipo')
    bristol_label = BRISTOL_LABEL.get(bristol_num, f'Tipo {bristol_num}') if bristol_num else '—'
    sintomas_act = seg.get('sintomasActivos', []) or []
    sintomas_str = ', '.join(sintomas_act) if sintomas_act else '—'

    vals10 = [
        fmt_fecha(seg.get('fecha') or seg.get('fechaCreacion')),
        v(seg, 'pesoActual'),
        bristol_label,
        v(seg, 'evacuacionesDia'),
        v(seg, 'toleranciaAlimentaria'),
        sintomas_str,
        v(seg, 'notasEvolucion'),
    ]
    for ci, val in enumerate(vals10):
        cell = row.cells[ci]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(str(val) if val else '—')
        r.font.size = Pt(9); r.font.name = 'Arial'; r.font.color.rgb = NEGRO
        set_cell_bg(cell, bg); set_cell_border(cell)

# ════════════════════════════════════════════════════════════════════════════
#  11. NOTAS CLÍNICAS Y OBSERVACIONES
# ════════════════════════════════════════════════════════════════════════════
heading_para('11. NOTAS CLÍNICAS Y OBSERVACIONES')

nota_gen = v(historia, 'notasAdicionales', default='')
if nota_gen and nota_gen != '—':
    p_nota = doc.add_paragraph()
    r_nota = p_nota.add_run(nota_gen)
    r_nota.font.size = Pt(10); r_nota.font.name = 'Arial'; r_nota.font.color.rgb = NEGRO

for _ in range(7):
    p_ln = doc.add_paragraph()
    p_ln.paragraph_format.space_before = Pt(14)
    p_ln.paragraph_format.space_after = Pt(0)
    pPr2 = p_ln._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    b2 = OxmlElement('w:bottom')
    b2.set(qn('w:val'), 'single'); b2.set(qn('w:sz'), '4')
    b2.set(qn('w:space'), '1'); b2.set(qn('w:color'), 'CCCCCC')
    pBdr2.append(b2); pPr2.append(pBdr2)

# Firma
p_firma = doc.add_paragraph()
p_firma.paragraph_format.space_before = Pt(30)
p_firma.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_f = p_firma.add_run('__________________________________\nLic. Karina Lara · Nutrióloga Pediátrica\nClínica Karina Lara')
r_f.font.size = Pt(10); r_f.font.name = 'Arial'; r_f.font.color.rgb = GRIS

# ─── Guardar ─────────────────────────────────────────────────────────────────
doc.save(output_path)
print(f"OK: {output_path}", flush=True)
